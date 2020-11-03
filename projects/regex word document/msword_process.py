import re
from docx import Document
import csv
import os

# This script takes in a word document, performs regular expressions to
# search for a pattern, and replaces the sentence with a pattern.

report = "Plains Midstream 2019 DRAFT Risk Report_r1.0 .docx"
report2 = "Plains Midstream 2019 DRAFT Risk Report_r1.0 (ALT) .docx"

##report = 'Plains Midstream 2019 DRAFT Risk Report Attachments_r1.0 .docx'
##report2 = 'Plains Midstream 2019 DRAFT Risk Report Attachments_r1.0 (ALT) .docx'

def docx_replace_regex(doc_obj, regex , replace):
    '''searches a document (doc_obj) and locates regex,
    and replaces it with replace.  doc_obj passed by reference
    so no need to return it.
    '''
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text, count=1 )
                    inline[i].text = text
                    return

doc = Document(report)
print('Loaded')

# 1 US gallon = 0.00378541 m3
regex = re.compile(r'([Ss]pill volume is between a minimum of )([0-9]{1,}\.[0-9]{1,})( and maximum of )([0-9]{1,}\.[0-9]{1,})( gallons)')
regex2 = re.compile(r'(Leak spill volume between a minimum of )([0-9]{1,}\.[0-9]{1,})( and maximum of )([0-9]{1,}\.[0-9]{1,})( gallons)')
regex3 = re.compile(r'(spill volume of )([0-9]{1,}\.[0-9]{1,})( gallons)')

for p in doc.paragraphs:
    if regex.search(p.text):
        text = list(regex.search(p.text).groups())
        text[1] = f"{float(text[1])*0.00378541:,.2f}"
        text[3] = f"{float(text[3])*0.00378541:,.2f}"
        text[4] = ' cubed meters'
        text = ''.join(text)
        text = regex.sub(text,p.text)
        p.text = text

for p in doc.paragraphs:
    if regex2.search(p.text):
        text = list(regex2.search(p.text).groups())
        text[1] = f"{float(text[1])*0.00378541:,.2f}"
        text[3] = f"{float(text[3])*0.00378541:,.2f}"
        text[4] = ' cubed meters'
        text = ''.join(text)
        text = regex2.sub(text,p.text)
        p.text = text
  
for p in doc.paragraphs:
    if regex3.search(p.text):
        text = list(regex3.search(p.text).groups())
        text[1] = f"{float(text[1])*0.00378541:,.2f}"
        text[2] = ' cubed meters'
        text = ''.join(text)
        text = regex3.sub(text,p.text)
        p.text = text
        
doc.save(report2)
print('Finished!')


#didn't seem to find this:
#   Leak spill volume between a minimum of 1691.87 and maximum of 2302.51 gallons
#   spill volume of 1976.71 gallons
